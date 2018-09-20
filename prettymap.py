#!/usr/bin/python
#   PrettyMap
# Author: 	phl4nk
# Date: 	07/05/2017
# Version:	v0.23

import os
import sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import itertools
from os import listdir
from os.path import isfile, join
from libnmap.parser import NmapParser,NmapParserException #EZMode
# Xtra Requirements
# sudo pip install python-libnmap
# http://libnmap.readthedocs.io for Ref
#sudo pip install python-docx
#https://python-docx.readthedocs.io/

# Comments/Notes, etc.
#
# Functionally Orientated for ease of modification
# trying to write 'generic' and 'high level' methods for potential later use
# NmapParser to begin with for EZ mode, currently just does XML output.
# Example Nmap command: nmap -sV -oX hostname.xml hostname
# NOTE: removeDuplicates() WILL DELETE DUPLICATE SCANS
#       SCANS WILL BE DELETED UNDER THESE CONDITIONS:
#           A) HOSTNAMES ARE THE SAME
#           B) IF NO HOSTNAMES, IF IP'S ARE THE SAME
#       SCANS __SHOULD__ BE AGGREGATED, BE WEARY OF DRAGONS!

# get files from current working directory
# if no dir is specified, use current working dir
# if no files are found, sys.exit
def gatherFiles(mydir=None):
    if not mydir:
        mydir = os.getcwd()
    print "[+] Searching",mydir,"for nmap XML files"
    #return full path of files (UNIX ONLY CURRENTLY)
    files = [mydir+'/'+f for f in listdir(mydir) if isfile(join(mydir, f))]
    if not files:
        print "[!] Error - no files to parse, Exiting"
        sys.exit(1)
    return files

# parse each individual nmap file with NmapParser, and turn it into a libnmap object
# returns a list of nmapObjects to be worked on
def parseFiles(fileList):
    nmapList = []
    for f in fileList:
        try:
            print "[+] Parsing", f
            nmapObject = NmapParser.parse_fromfile(f)
            nmapList.append(nmapObject)
        except NmapParserException,e:
            print "[!] Nmap Parsing Exception:",e,f
            # Do something clever? no, just ignore it
        except Exception,e:
            print "[!] Generic Exception:",e
            # Again, Ostrich algorithm
    if not nmapList:
        print "[!] Error - no nmap scans detected Exiting"
        sys.exit(1)
    return nmapList


# some crazy ass innificient looping to try and reduce duplicates
# iterates through all the parsed nmap files, and compares each host
# if hostname is the same, its conisdered a duplicate
# if no hostname, the checks if IP is the same
# any missing services should be added to the host for aggregation
# this should theoretically combine UDP and TCP scans into one host
# I sacrificed a lizard, chicken, and a small child to get this working...
# TODO: update scan with banner if has one (check for changed())
def removeDuplicates(nmapObjectList):
    # Yo Dawg, i heard you like loops?
    for x,y in itertools.combinations(nmapObjectList, 2):
        for hostsX in x.hosts:
            for hostsY in y.hosts:
                #hostname dupe check:
                for hostnameX in hostsX.hostnames:
                    if hostnameX in hostsY.hostnames:
                        print "[+] Duplicate hostname found",hostsY
                        #is there missing services?
                        addMissingServices(hostsX,hostsY)
                        print "[+] Deleting scan",hostsY
                        y.hosts.remove(hostsY)
                        break
                # no assicated hostname, so check if any IP's are the same.
                if not hostsX.hostnames:
                    if hostsX.ipv4 == hostsY.ipv4:
                        print "[+] Duplicate IP found",hostsY
                        addMissingServices(hostsX,hostsY)
                        print "[+] Deleting scan",hostsY
                        y.hosts.remove(hostsY)
                        break
    return nmapObjectList

# compares two nmap scans
# if nmap2 has any services nmap 1 doesn't, add them
def addMissingServices(nmap1,nmap2):
    removed = nmap1.diff(nmap2).removed()
    if removed:
        print "[+] Missing data, adding services"
        for service2 in nmap2.services:
            if service2 not in nmap1.services:
                print "[+] Adding service",service2
                nmap1.services.append(service2)

# Create the table ~ super messy, need to clean this up
def generateTable(listofNmaps):
    print "[+] Generating new document...such excite"
    document = Document()
    for nmap in listofNmaps:
        for host in nmap.hosts:
            print "[+] Creating entry for:",host
            # create header of IP - hostname1,hostname2,etc.
            ip = host.ipv4
            hostname = ', '.join(host.hostnames)
            header = ip+" - "+hostname if hostname else ip
            document.add_heading(header, level=1)
            document.add_paragraph("\n")
            if not host.services:
                document.add_paragraph("No services detected")
                break
            # TCP Table
            table = document.add_table(rows=1, cols=4)
            hdrtcp_cells = table.rows[0].cells
            tcpheader = hdrtcp_cells[0].add_paragraph('TCP')
            tcpheader.alignment=WD_ALIGN_PARAGRAPH.CENTER
            hdrtcp_cells[1].text = 'Port'
            hdrtcp_cells[2].text = 'State'
            hdrtcp_cells[3].text = 'Banner'
            #can only insert row at bottom, so need to store UDP's
            udpServices = []
            for service in host.services:
                if service.protocol == 'udp':
                    udpServices.append(service)
                    pass
                row_cells = table.add_row().cells
                row_cells[0].merge(hdrtcp_cells[0])
                row_cells[1].text=str(service.port)
                row_cells[2].text=service.state
                row_cells[3].text=service.banner
            #UDP Table
            firstRun=True
            for service in udpServices:
                if firstRun:
                    root_cell = table.add_row().cells
                    udpheader = root_cell[0].add_paragraph('UDP')
                    udpheader.alignment=WD_ALIGN_PARAGRAPH.CENTER
                    root_cell[1].text=str(service.port)
                    root_cell[2].text=service.state
                    root_cell[3].text=service.banner
                else:
                    row_cells = table.add_row().cells
                    row_cells[0].merge(root_cell[0])
                    row_cells[1].text=str(service.port)
                    row_cells[2].text=service.state
                    row_cells[3].text=service.banner
                firstRun=False
    document.save('demo.docx')


fileList = gatherFiles()
nmapList = parseFiles(fileList)
newList = removeDuplicates(nmapList)
generateTable(newList)

# for n in newList:
#    for hosts in n.hosts:
#        print hosts
#        for service in hosts.services:
#            print "\t",service
# sys.exit(1)
